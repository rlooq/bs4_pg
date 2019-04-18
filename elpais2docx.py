# This is a script to grab text from El País articles, taking the headline, subtitle and ´
# main article, and saving it to a docx.

import bs4
from bs4 import BeautifulSoup as soup
import requests
from requests.exceptions import HTTPError

from datetime import datetime
from docx import Document

my_url=""
# Adding a browser user agent
user_agent = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:65.0) Gecko/20100101 Firefox/65.0'
headers = {'User-Agent': user_agent}

try:
    response = requests.get(my_url, headers=headers)
    response.raise_for_status()
    # If the response is successful, no exception will be raised
    page=soup(response.text, "html.parser")
    document = Document("article_template.docx")
    # Extracting relevant text
    title=page.h1.text
    subtitle=page.h2.text
    author=page.find("span", {"class": "autor-nombre"})
    document.add_heading(title)
    document.add_heading(subtitle, level=2)
    document.add_paragraph(author.text)
    
    for par in page.findAll('p'):
        document.add_paragraph(par.text)
    
    today=datetime.now().strftime("%Y-%m-%d_%H-%M")
    document.add_paragraph(today)
    document.save('scraped_elpais_{}.docx'.format(today))

except HTTPError as http_err:
    print('HTTP error occurred: {}'.format(http_err))
except Exception as err:
    print('Other error occurred: {}'.format(err))
else:
        print('Success!')